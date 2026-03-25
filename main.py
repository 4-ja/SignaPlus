import tkinter as tk
from tkinterdnd2 import DND_FILES, TkinterDnD
from PIL import Image
import pystray
from pystray import MenuItem as item
import threading
import ctypes
import os
import time
import shutil
import winsound
import configparser
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document 
from docx.shared import Inches

import utils
import gui

# Force Windows to show the app icon on the Taskbar
myappid = 'com.autosign.ultimate.v2.6' 
try:
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
except:
    pass

class AppController:
    def __init__(self):
        self.root = TkinterDnD.Tk()
        self.root.withdraw() 
        
        self.icon_file = "icon.ico"
        
        # Set the icon for the hidden root to help taskbar grouping
        if os.path.exists(self.icon_file):
            self.root.iconbitmap(self.icon_file)
        
        utils.ensure_settings_exists()
        utils.create_desktop_shortcut()
        
        self.load_config()
        
        self.gui_window = None
        self.start_watcher()
        self.setup_tray()

    def load_config(self):
        """Loads or reloads settings from the text file."""
        self.config = configparser.ConfigParser()
        self.config.read('settings.txt')
        self.target_name = self.config.get('SETTINGS', 'SignatoryName', fallback='John Doe')

    def open_settings(self):
        """Triggered by the gear icon in the GUI."""
        if self.gui_window:
            gui.SettingsGUI(self.gui_window, self.target_name, self.save_settings)

    def open_help(self):
        """Triggered by the question mark icon in the GUI."""
        if self.gui_window:
            gui.HelpGUI(self.gui_window)

    def save_settings(self, new_name):
        """Callback from settings window to update app state and file."""
        self.target_name = new_name
        if not self.config.has_section('SETTINGS'):
            self.config.add_section('SETTINGS')
        self.config.set('SETTINGS', 'SignatoryName', new_name)
        with open('settings.txt', 'w') as configfile:
            self.config.write(configfile)
        print(f"Settings saved: {new_name}")

    def get_tray_icon(self):
        try:
            if os.path.exists(self.icon_file):
                return Image.open(self.icon_file)
        except:
            pass
        return Image.new('RGB', (64, 64), (0, 120, 215))

    def show_window(self):
        if not self.gui_window or not tk.Toplevel.winfo_exists(self.gui_window):
            # Pass both the open_settings and open_help functions as callbacks
            self.gui_window = gui.DownloadStyleGUI(self.root, self.open_settings, self.open_help)
            
            if os.path.exists(self.icon_file):
                self.gui_window.iconbitmap(self.icon_file)
            
            self.gui_window.drop_target_register(DND_FILES)
            self.gui_window.dnd_bind('<<Drop>>', self.handle_drop)
            
            # Binding the double-click to reveal the file
            if hasattr(self.gui_window, 'tree'):
                self.gui_window.tree.bind("<Double-1>", self.on_item_double_click)
            
            # Load the existing files from the folder
            self.load_history_to_gui()
        else:
            self.gui_window.deiconify()
            self.gui_window.lift()

    def load_history_to_gui(self):
        base_dir = 'Signed_Documents'
        if not os.path.exists(base_dir): return

        for item in self.gui_window.tree.get_children():
            self.gui_window.tree.delete(item)

        all_signed_files = []
        for root, dirs, files in os.walk(base_dir):
            for file in files:
                file_path = os.path.join(root, file)
                mtime = os.path.getmtime(file_path)
                date_str = datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M')
                all_signed_files.append((file, date_str, mtime))

        all_signed_files.sort(key=lambda x: x[2], reverse=True)

        for file_name, date_str, _ in all_signed_files:
            self.gui_window.add_task_row(file_name, date_str, "Completed")

    def on_item_double_click(self, event):
        try:
            tree = self.gui_window.tree
            selection = tree.selection()
            if not selection: return
            file_name = tree.item(selection[0], "values")[0]
            
            base_dir = 'Signed_Documents'
            for root, dirs, files in os.walk(base_dir):
                if file_name in files:
                    full_path = os.path.join(os.path.abspath(root), file_name)
                    os.system(f'explorer /select,"{full_path}"')
                    return
        except Exception as e:
            print(f"Open Error: {e}")

    def handle_drop(self, event):
        files = self.root.tk.splitlist(event.data)
        for f_path in files:
            if os.path.isfile(f_path):
                threading.Thread(target=self.process_file_engine, args=(f_path, True), daemon=True).start()

    def process_file_engine(self, src_path, is_manual_drop=False):
        base_output_dir = 'Signed_Documents'
        sig_folder = 'Signature'
        now = datetime.now()
        today_folder = now.strftime('%Y-%m-%d')
        date_display = now.strftime('%Y-%m-%d %H:%M')
        
        final_output_dir = os.path.join(base_output_dir, today_folder)
        if not os.path.exists(final_output_dir): 
            os.makedirs(final_output_dir)

        file_name = os.path.basename(src_path)
        name_part, ext = os.path.splitext(file_name)
        ext = ext.lower()
        new_file_name = f"{name_part}_signed{ext}"
        dest_path = os.path.join(final_output_dir, new_file_name)

        sig_path = None
        if os.path.exists(sig_folder):
            for f in os.listdir(sig_folder):
                if f.lower().startswith("signature"):
                    sig_path = os.path.join(sig_folder, f)
                    break

        success = False
        if sig_path:
            try:
                if ext == '.pdf':
                    doc = fitz.open(src_path)
                    found_text = False
                    for page_index in range(len(doc) - 1, -1, -1):
                        page = doc[page_index]
                        matches = page.search_for(self.target_name)
                        if matches:
                            inst = matches[-1]
                            sig_w, sig_h = 120, 60
                            x0 = inst.x0 + (inst.width / 2) - (sig_w / 2)
                            y0 = inst.y0 - sig_h + 15 
                            page.insert_image(fitz.Rect(x0, y0, x0 + sig_w, y0 + sig_h), filename=sig_path)
                            found_text = True
                            break
                    if not found_text:
                        last_page = doc[-1]
                        rect = fitz.Rect(last_page.rect.width - 150, last_page.rect.height - 150, 
                                          last_page.rect.width - 50, last_page.rect.height - 50)
                        last_page.insert_image(rect, filename=sig_path)
                    doc.save(dest_path)
                    doc.close()
                    success = True
                elif ext in ['.docx', '.doc']:
                    doc = Document(src_path)
                    doc.add_paragraph(f"\nElectronically Signed for: {self.target_name}")
                    doc.add_picture(sig_path, width=Inches(1.5))
                    doc.save(dest_path)
                    success = True
                elif ext in ['.png', '.jpg', '.jpeg']:
                    img = Image.open(src_path).convert("RGBA")
                    sig = Image.open(sig_path).convert("RGBA")
                    scale = int(img.width * 0.20)
                    sig.thumbnail((scale, scale))
                    img.paste(sig, (img.width - sig.width - 50, img.height - sig.height - 50), sig)
                    img.convert("RGB").save(dest_path)
                    success = True
                
                if success:
                    winsound.PlaySound("SystemAsterisk", winsound.SND_ALIAS)
                    if not is_manual_drop: os.remove(src_path)
            except Exception as e:
                print(f"Engine Error: {e}")
                if not is_manual_drop: shutil.move(src_path, dest_path)
        else:
            if not is_manual_drop: shutil.move(src_path, dest_path)

        if self.gui_window and tk.Toplevel.winfo_exists(self.gui_window):
            status = "Completed" if success else "Failed"
            self.root.after(0, self.gui_window.add_task_row, new_file_name, date_display, status)

    def start_watcher(self):
        def watch_loop():
            input_dir = 'To_Be_Signed'
            while True:
                try:
                    if not os.path.exists(input_dir): os.makedirs(input_dir)
                    files = [f for f in os.listdir(input_dir) if os.path.isfile(os.path.join(input_dir, f))]
                    for file_name in files:
                        self.process_file_engine(os.path.join(input_dir, file_name), is_manual_drop=False)
                except:
                    pass
                time.sleep(3)
        threading.Thread(target=watch_loop, daemon=True).start()

    def quit_app(self, icon):
        icon.stop()
        self.root.quit()

    def setup_tray(self):
        menu = (item('Open Dashboard', self.show_window), item('Exit', self.quit_app))
        self.icon = pystray.Icon("AutoSign", self.get_tray_icon(), "AutoSignatory", menu)
        threading.Thread(target=self.icon.run, daemon=True).start()

if __name__ == "__main__":
    app = AppController()
    app.show_window() 
    app.root.mainloop()